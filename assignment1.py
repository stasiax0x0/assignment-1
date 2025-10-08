from collections import defaultdict     #for creating a dictionary (with default default values) of lists
from datetime import datetime           #for parsing timestamps
from datetime import timedelta          #for working with time differences
import docx                             #for creating word documents
from docx import Document               #specifically for Document class
from docx.shared import Inches          #for setting image sizes in Word
from docx.shared import RGBColor        #for setting text colours in word
from rich import print                      #pretty printing
from rich.console import Console           #for pretty printing to console
from rich.table import Table             #for pretty printing tables
import typer #importing typer library which helps create cli applications

LOGFILE= "CA1_project.log"      #defining log file name that we are analyzing
#1.	Parse log files line by line (files provided).

def parse_auth_line(line):              #function to parse a line from the auth log
    
    parts = line.split()                #split the line into parts
    # timestamp: first 3 tokens 'Mar 10 13:58:01'
    ts_str = " ".join(parts[0:3])
    try:
        ts = datetime.strptime(f"2025 {ts_str}", "%Y %b %d %H:%M:%S")       #parse the timestamp, assuming year 2025
    except Exception:                                                       #if parsing fails, set timestamp to None
        ts = None
    ip = None                                                               #initialize ip to None
    event_type = "other"                                                    #initialize event_type to "other"
if "Failed password" in line:                                               #Check what type of login event this is 
        event_type = "failed"                                               #mark as failed login attempt
    elif "Accepted password" in line or "Accepted publickey" in line:
        event_type = "accepted"                                             #mark as successful login
    if " from " in line:                                                    #look for the "from" token to find the IP address
        try:
            idx = parts.index("from")                                       #find position of "from" in tokens 
            ip = parts[idx+1]                                               #IP address should be the next token
        except (ValueError, IndexError):                                    #if "from" not found or no token after it, set ip to None#
            ip = None                                                       #set IP to none
    return ts, ip, event_type                                               #return extracted info


#2.  Detect and count failed login attempts, grouping by source IP.
#main script

if __name__ == "__main__":
    #create a dictionary to store timestamps for each IP
    counts= defaultdict(int)                                #count failed attemps per IP with default value at 0
    per_ip_timestamps = defaultdict(list)                   #store timestamps for each IP with a empty list as default    
    with open(LOGFILE) as f:                                #open the log file
        for line in f:                                      #read each line
            ts, ip, event = parse_auth_line(line)           #parse the line to get timestamp, ip, and event type
            if ip and event == "failed":                    #checks that ip is not null, and that event=="failed"
                counts[ip] += 1                             #imcrement count for this ip
            if ts:                                          #checks that ts is not null
                per_ip_timestamps[ip].append(ts)            #add the timestamp to the list for this IP
    console = Console()                                     #create a console object for pretty printing
    console.print("[bold red]Failed login attempts: [/bold red]", dict(counts))  #pretty print the counts dictionary in bold red

print("")                                                   #print empty line for space

#3.	Identify possible brute-force attacks (≥ 5 failed logins from one IP within 10 minutes).

incidents = []                                                          #list to store detected brute-force incidents
window = timedelta(minutes=10)                                          #define the time window for detecting brute-force attempts (10 minutes)
for ip, times in per_ip_timestamps.items():                             #for each IP and its list of timestamps
    if not ip:
        continue                                                        #skip entry if the IP is None
    times.sort()                                                        #sort the timestamps
    n = len(times)                                                      #get the number of timestamps
    i = 0                                                               #initialize the start index
    while i < n:                                                        #scan through the timestamps for this IP
        j = i                                                           #initialize the end index
        while j + 1 < n and (times[j+1] - times[i]) <= window:          #expand the window: as long as the next timestamp is within 10 minutes of the first one in the current window
            j += 1                                                      #move the end index forward
        count = j - i + 1                                               #calculate the number of failed attempts in this window
        if count >= 5:                                                  #if there are 5 or more failed attempts in this window, record it as an incident
            incidents.append({                                          #add the incident details to the list
                "ip": ip,
                "count": count,
                "first": times[i].isoformat(),                          #convert to string
                "last": times[j].isoformat()
            })
            # advance i past this cluster to avoid duplicate overlapping reports:
            i = j + 1
        else:                                                            #move to the next timestamp if no incident detected
            i += 1
# print incidents
console = Console()
console.print(f'[bold red]Detected {len(incidents)} brute-force incidents:[/bold red]')  #pretty print the number of detected incidents in bold red
table = Table(title="Brute-force Incidents")               #create a table to display the incidents
table.add_column("IP", style="cyan", no_wrap=True)         #add IP column
table.add_column("Failed Attempts", style="magenta")       #add count column (magenta)
table.add_column("Time Window", style="green")             #add time range column (green)
for incident in incidents:                                 #add a row to the table for each incident
    table.add_row(
    incident["ip"],                                        #IP address
    str(incident["count"]),                                #number of attempts
        f'{incident["first"]} to {incident["last"]}'       #time range
    )
console.print(table)                                      #pretty print the table to the console
print(" ")                                                #Empty line for spacing  




#ADDITIONAL FEATURES
#7.	Visualise findings (bar chart of attacker IPs).
# --- Optional: Plot bar chart of tot incidents per ip ---
import matplotlib.pyplot as plt     # import plotting library


# get the ips and counts
allips = sorted(counts.items(), key=lambda x: x[1], reverse=True)[:16]              # get the 16 IPs with failed attempts
ips = [ip for ip, _ in allips]                                                      # list of IPs for the chart, underscore _ means we ignore the second value
top_counts = [count for _, count in allips]                                         # list of counts for the chart

plt.figure(figsize=(16,5))                                                          # set size of chart
plt.bar(ips, top_counts, color=["#FFB6C1", "#FF69B4"])                              # create bar chart and sets the colors
plt.title("Total failed login attempts per IP", fontweight='bold', fontsize=14 )    # set title and labels
plt.xlabel("IP", fontweight='bold', fontsize=12)                                    # set x-axis label
plt.ylabel("Failed attempts", fontweight='bold', fontsize=12)                       # set y-axis label
plt.xticks(rotation=45)                                                             # rotate ip labels for readability
plt.tight_layout()                                                                  # adjust layout to fit everything
plt.savefig("failed_logins.png")                                                    # save chart as image file
plt.show()                                                                          # display the chart

console= Console()
console.print("[bold red]Bar chart saved to failed_logins.png[/bold red]")          # notify user chart is saved
print(" ")

#4.	Output results into a structured report.


doc = Document()                                                        #create new paragraph
doc.add_heading("BRUTE FORCE INCIDENTS REPORT", 0)                      # add title to report
                    #make it in bold
p = doc.add_paragraph()
run = p.add_run("Total failed login attempts per IP:")                  #add text
run.bold = True                                                         #make text bold

for ip, count in counts.items():                                        #add each IP and its counts to document
    doc.add_paragraph(f"{ip}:   {count} failed attempts")


#add brute-force incidents summary   
p = doc.add_paragraph()                                                  #new paragraph
run = p.add_run("Detected {} brute-force incidents".format(len(incidents))) #add text with count
run.bold = True                                                              #make the number of incidents bold
run.font.color.rgb = RGBColor(255, 0, 0)                                     # RGB for red
doc.add_paragraph(" \n\n")                                                     #add empty lines


#add bar chart to docuement
p = doc.add_paragraph()                                                  #new paragraph
run = p.add_run("Total failed login attempts per IP BAR CHART:")        #add text
run.bold = True                                                         #make bold 
doc.add_picture('failed_logins.png', width=Inches(7))                   #add saved chart image


#add details fro each detected incident
for i, incident in enumerate(incidents, 1):                             #number incidents starting from 1
    p = doc.add_paragraph()                                             #new paragraph
    run = p.add_run(f'Incident {i}:')                                   #add incident header
    run.bold = True                                                     #make bold
    doc.add_paragraph(f'IP: {incident["ip"]}')                          #add ip 
    doc.add_paragraph(f'Failed Attempts: {incident["count"]}')          #add attempt count
    doc.add_paragraph(f'Time Window: {incident["first"]} to {incident["last"]}\n\n')    #add time ramge 

doc.save('report.docx')                                                 #save to word doc



#ADDITIONAL FEATURE 5.	Implement as a command-line application (using Typer or Click).

app = typer.Typer()                                                         # create a typer application instance - this is main objecy for the CLI app

@app.command()                                                              # decorator tells typer that function below should be a cli command 
def analyze_ips(logfile: str = "CA1_project.log"):                          # define function called "count_ips" that takes ine parameter, looks thru CA1_project.log file
    all_unique_ips = set()                                                  #creates empty set "unique_ips" to store unique ips, set automatically removed applications
    unique_failed_ips = set()                                               #creates empty set "unique_failed_ips" to store unique ips with failed logins
    total_failed_attempts = 0                                               #counter for total failed login attempts
    with open(logfile) as f:                                                #opens  logfile for reading, "with open" automatically closes file
        for line in f:                                                      # loops thru each line in file one by one 
            if "from " in line:                                             #checks if current line contains text "from"
                parts = line.split()                                        #splits line into separae words: creates list of words
                try:                                                         # try to find and extract ip adress (might fail, so use try/except)
                    idx = parts.index("from")                               #find position of word from in parts line 
                    ip = parts[idx+1]                                       #gets word immediately after "from" (ip address")
                    all_unique_ips.add(ip)                                  #adds ip to set (automatically handles duplicates)
                    if "Failed password" in line:                           #checks if line contains "Failed password"
                
                        unique_failed_ips.add(ip)                           #adds ip to set of failed ips
                        total_failed_attempts += 1                          #increments counter of total failed attempts
                except (ValueError, IndexError):                            #if smth wrong (like "from" not found or no ip after it) skip line  
                    continue                                                #skips to next line on file 
    typer.echo(f"Overall unique IPs: {len(all_unique_ips)}")                    #shows count all unique ips founf in log
    typer.echo(f"Unique IPs with failed logins: {len(unique_failed_ips)}")      #shows count of unique ips with failed logins
    typer.echo(f"Total failed login attempts: {total_failed_attempts}")         #shows total number of failed login attempts

    if all_unique_ips:                                                          # avoids division by zero if no ips found
        percentage = (len(unique_failed_ips) / len(all_unique_ips)) * 100                         #calculates percentage of unique ips with failed logins
        typer.echo(f"Percentage of unique IPs with failed logins: {percentage:.2f}%")             #shows percentage formatted to 2 decimal places
if __name__ == "__main__":                                                                        #special condition checks if script is being run direactly (not imprted)
    app()                                                                                          # if running directly, start typer application

                                